import streamlit as st
from docx import Document
import fitz  # PyMuPDF
import os

def extract_images_from_page(page):
    images = []
    for img_index, img in enumerate(page.get_images(full=True)):
        xref = img[0]
        base_image = pdf_document.extract_image(xref)
        image_bytes = base_image["image"]
        images.append(image_bytes)
    return images

def convert_pdf_to_word(pdf_path, word_path):
    # 打开PDF文件
    pdf_document = fitz.open(pdf_path)
    # 创建一个新的Word文档
    doc = Document()

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")
        images = extract_images_from_page(page)
        
        # 添加文本到Word文档
        doc.add_paragraph(text)

        # 添加图片到Word文档
        for image in images:
            doc.add_picture(image, width=docx.shared.Inches(3))  # 插入图片（示例）

    # 保存Word文档
    doc.save(word_path)

# 主函数
def app1():
    st.title("PDF 转 Word 批量转换工具")
    
    # 上传文件
    uploaded_files = st.file_uploader("选择一个或多个PDF文件", type="pdf", accept_multiple_files=True)
    
    if uploaded_files:
        output_dir = "output"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        for uploaded_file in uploaded_files:
            pdf_filename = uploaded_file.name
            word_filename = pdf_filename.replace('.pdf', '.docx')
            pdf_path = os.path.join(output_dir, pdf_filename)
            word_path = os.path.join(output_dir, word_filename)
            
            with open(pdf_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            convert_pdf_to_word(pdf_path, word_path)
            
            st.success(f"转换成功: {word_filename}")
            with open(word_path, "rb") as f:
                st.download_button(f"下载 {word_filename}", f, file_name=word_filename)

if __name__ == "__main__":
    app1()
